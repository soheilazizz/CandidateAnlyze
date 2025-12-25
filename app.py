import os
import io
import re
import html
import base64
from datetime import date
from pathlib import Path
from urllib.parse import urlparse

import streamlit as st
from openai import OpenAI

# =========================
# Page config + icon
# =========================
ICON_PATH = "assets/icon.png"
if Path(ICON_PATH).exists():
    st.set_page_config(page_title="ارزیابی کاندیدا", page_icon=ICON_PATH, layout="wide")
else:
    st.set_page_config(page_title="ارزیابی کاندیدا", page_icon="🧾", layout="wide")

# =========================
# Font embedding (Vazirmatn) + RTL CSS
# =========================
def load_font_b64(path: str) -> str:
    return base64.b64encode(Path(path).read_bytes()).decode()

css_font_faces = ""
reg_path = Path("assets/Vazirmatn-Regular.ttf")
bold_path = Path("assets/Vazirmatn-Bold.ttf")

if reg_path.exists() and bold_path.exists():
    REG = load_font_b64(str(reg_path))
    BOLD = load_font_b64(str(bold_path))
    css_font_faces = f"""
    @font-face {{
      font-family: 'FA';
      src: url(data:font/ttf;base64,{REG}) format('truetype');
      font-weight: 400;
      font-style: normal;
    }}
    @font-face {{
      font-family: 'FA';
      src: url(data:font/ttf;base64,{BOLD}) format('truetype');
      font-weight: 700;
      font-style: normal;
    }}
    """

st.markdown(f"""
<style>
{css_font_faces}

:root {{
  --bg: #0b1220;
  --card: rgba(255,255,255,0.06);
  --border: rgba(255,255,255,0.12);
  --text: rgba(255,255,255,0.93);
  --muted: rgba(255,255,255,0.75);
}}

html, body, [class*="css"] {{
  font-family: {'FA' if css_font_faces else 'system-ui'}, -apple-system, Segoe UI, Roboto, Arial, sans-serif !important;
  direction: rtl;
  text-align: right;
  line-height: 2.05;
  font-size: 15.5px;
  color: var(--text);
  unicode-bidi: plaintext;
}}

h1,h2,h3 {{
  letter-spacing: 0 !important;
  line-height: 1.5;
}}

.card {{
  background: var(--card);
  border: 1px solid var(--border);
  padding: 18px 18px;
  border-radius: 16px;
}}

.small {{
  font-size: 0.92rem;
  color: var(--muted);
}}

hr {{
  border: none;
  height: 1px;
  background: var(--border);
  margin: 16px 0;
}}

.stButton>button {{
  border-radius: 12px;
  padding: 10px 16px;
  font-weight: 700;
}}

.report {{
  direction: rtl;
  text-align: right;
  unicode-bidi: plaintext;
  line-height: 2.05;
  font-size: 15.5px;
}}

.report p {{
  margin: 0 0 10px 0;
}}

.report table {{
  width: 100%;
  border-collapse: collapse !important;
  margin: 10px 0 16px 0;
}}

.report th, .report td {{
  border: 1px solid rgba(255,255,255,0.28) !important;
  padding: 10px 10px !important;
  vertical-align: top !important;
  text-align: right !important;
}}

.report thead th {{
  background: rgba(255,255,255,0.08);
  font-weight: 700;
}}

.ltr {{
  direction: ltr;
  unicode-bidi: embed;
  display: inline-block;
  text-align: left;
}}
</style>
""", unsafe_allow_html=True)

# =========================
# Helpers: Markdown table parsing
# =========================
def is_md_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 3

def is_md_separator(line: str) -> bool:
    s = line.strip().replace(" ", "")
    return s.startswith("|") and set(s.replace("|", "")) <= set("-:")

def parse_md_table(lines):
    header = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    rows = []
    for ln in lines[2:]:
        cols = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cols) < len(header):
            cols += [""] * (len(header) - len(cols))
        rows.append(cols[:len(header)])
    return header, rows

latin_chunk = re.compile(r'([A-Za-z0-9][A-Za-z0-9\-\._/+# ]{0,50})')

def wrap_ltr(text: str) -> str:
    safe = html.escape(text)
    return latin_chunk.sub(r'<span class="ltr">\1</span>', safe)

def markdown_to_html_with_tables(md: str) -> str:
    lines = md.splitlines()
    out = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        s = line.strip()

        # table block
        if is_md_table_line(s) and i + 1 < len(lines) and is_md_separator(lines[i+1]):
            block = [s, lines[i+1].strip()]
            j = i + 2
            while j < len(lines) and is_md_table_line(lines[j].strip()):
                block.append(lines[j].strip())
                j += 1

            headers, rows = parse_md_table(block)
            out.append("<table><thead><tr>" + "".join(f"<th>{wrap_ltr(h)}</th>" for h in headers) + "</tr></thead><tbody>")
            for r in rows:
                out.append("<tr>" + "".join(f"<td>{wrap_ltr(c)}</td>" for c in r) + "</tr>")
            out.append("</tbody></table>")
            i = j
            continue

        # normal text
        if s:
            out.append(f"<p>{wrap_ltr(s)}</p>")
        else:
            out.append("<div style='height:6px'></div>")
        i += 1

    return "<div class='report'>" + "\n".join(out) + "</div>"

# =========================
# Helpers: Extract resume/JD text
# =========================
def extract_text_from_upload(uploaded_file) -> str:
    name = uploaded_file.name.lower()

    if name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)

    if name.endswith(".docx"):
        from docx import Document
        tmp = io.BytesIO(uploaded_file.getvalue())
        doc = Document(tmp)
        return "\n".join([p.text for p in doc.paragraphs])

    raise ValueError("فرمت رزومه/JD باید یکی از pdf/docx/txt باشد.")

# =========================
# AvalAI calls
# =========================
def transcribe_audio_bytes(file_bytes: bytes, filename: str) -> str:
    api_key = os.getenv("AVALAI_API_KEY")
    if not api_key:
        raise RuntimeError("کلید AVALAI_API_KEY در Secrets ست نشده.")

    client = OpenAI(base_url="https://api.avalai.ir/v1", api_key=api_key)

    tmp_path = Path("/tmp") / filename
    tmp_path.write_bytes(file_bytes)

    try:
        with open(tmp_path, "rb") as f:
            t = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="text",
                language="fa",
            )
        return str(t)
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

def generate_report(resume_text: str, jd_text: str, interview_asr: str) -> str:
    api_key = os.getenv("AVALAI_API_KEY")
    if not api_key:
        raise RuntimeError("کلید AVALAI_API_KEY در Secrets ست نشده.")
    client = OpenAI(base_url="https://api.avalai.ir/v1", api_key=api_key)

    SYSTEM = "You are a strict, evidence-based HR evaluator. Be structured, concise, and avoid fluff."

    ASR_NOTE_FA = """
متن مصاحبه زیر خروجی خام تبدیل گفتار به متن (ASR) است و ممکن است شامل غلط املایی،
شکست کلمات، تکرار، یا خطاهای نگارشی باشد که ناشی از سیستم ASR است نه فرد مصاحبه‌شونده.
کیفیت زبان/املا را معیار قضاوت قرار نده. خطاهای متنی را نویز فرض کن.
"""

    FORMAT_SPEC = f"""
گزارش ارزیابی کاندیدا — نسخه یک‌صفحه‌ای (فارسی)

نام کاندیدا: (از رزومه استخراج کن؛ اگر نبود "نامشخص")
عنوان شغل: (از JD استخراج کن؛ اگر نبود "نامشخص")
تاریخ گزارش: {date.today().strftime("%Y-%m-%d")}
منابع بررسی: رزومه + فایل صوتی مصاحبه (خروجی ASR)

1) جمع‌بندی مدیریتی
- امتیاز تناسب کلی (Fit Score): XX/100 | سطح اطمینان: کم/متوسط/بالا
- پیشنهاد: Yes / No / Maybe (در صورت نیاز مشروط)
- چرا مثبت؟ (۲-۳ جمله)
- ریسک اصلی: (۱-۲ جمله)

2) نقاط قوت کلیدی (Strengths)
3) نقاط ضعف / ریسک‌ها (Weaknesses & Risks)

4) تحلیل تناسب مهارتی (Resume vs JD)
- Must-have ها (کلیدی): جدول Markdown دقیق با 3 ستون:
| نیاز شغلی | شواهد از رزومه/مصاحبه | میزان تطابق |
|---|---|---|
| ... | ... | پایین/متوسط/بالا |

- شکاف‌ها (Gaps): 2 تا 4 مورد با Impact: Low/Medium/High

5) تحلیل مصاحبه (سبک کاری از لحن و پاسخ‌ها)
- ساختار/شفافیت/مالکیت/ریسک‌های ارتباطی + دلیل کوتاه
- 2 Quote کوتاه

6) سازگاری ادعاها (Resume vs Interview)
7) پیشنهاد برای دور بعد (سوالات هدفمند)
8) نتیجه نهایی

مبنای Fit Score را صریح و ساده توضیح بده:
- 4 معیار: درک استراتژیک، تحلیل و تصمیم‌گیری، نگاه اجرایی، نشانه‌های رفتاری
- برای هر معیار: امتیاز + شواهد + دلیل
"""

    prompt = f"""
{FORMAT_SPEC}

قوانین:
- فقط بر اساس سه ورودی قضاوت کن: JD، رزومه، متن مصاحبه
- متن مصاحبه ASR خام است: {ASR_NOTE_FA}
- اگر داده نداریم: "نامشخص/یافت نشد"
- از قطعیت‌نمایی روانشناسانه پرهیز کن؛ نشانه‌ها را احتمالی بیان کن.

[JD]
{jd_text}

[RESUME]
{resume_text}

[INTERVIEW_ASR]
{interview_asr}
"""

    resp = client.chat.completions.create(
        model=os.getenv("AVALAI_TEXT_MODEL", "gpt-4o-mini"),
        messages=[
            {"role": "system", "content": SYSTEM},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return resp.choices[0].message.content

# =========================
# Word export: convert ALL markdown tables to real Word tables
# =========================
def report_to_docx_bytes(report_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    def add_rtl_paragraph(text, bold=False, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    add_rtl_paragraph("گزارش ارزیابی کاندیدا — نسخه یک‌صفحه‌ای", bold=True, size=14)
    doc.add_paragraph("")

    lines = report_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph("")
            i += 1
            continue

        # detect markdown table start
        if is_md_table_line(line) and i + 1 < len(lines) and is_md_separator(lines[i+1]):
            table_block = [line, lines[i+1].strip()]
            j = i + 2
            while j < len(lines) and is_md_table_line(lines[j].strip()):
                table_block.append(lines[j].strip())
                j += 1

            headers, rows = parse_md_table(table_block)

            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"

            # header row
            hdr_cells = t.rows[0].cells
            for c_idx, h in enumerate(headers):
                hdr_cells[c_idx].text = h

            # rows
            for r in rows:
                row_cells = t.add_row().cells
                for c_idx, val in enumerate(r):
                    row_cells[c_idx].text = val

            # align all cells right and set font size
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        for run in p.runs:
                            run.font.size = Pt(10)

            doc.add_paragraph("")
            i = j
            continue

        # headings/meta
        is_section = any(line.startswith(f"{k})") for k in range(1, 9))
        is_meta = line.startswith("نام کاندیدا") or line.startswith("عنوان شغل") or line.startswith("تاریخ گزارش") or line.startswith("منابع بررسی")

        if is_section:
            add_rtl_paragraph(line, bold=True, size=12)
        elif is_meta:
            add_rtl_paragraph(line, bold=True, size=11)
        else:
            add_rtl_paragraph(line, bold=False, size=11)

        i += 1

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =========================
# UI
# =========================
st.markdown("""
<div class="rtl">
  <h1>🧠 ارزیابی کاندیدا</h1>
  <div class="small">آپلود ۳ فایل → تبدیل گفتار به متن → گزارش ساختاریافته → خروجی Word</div>
</div>
<hr/>
""", unsafe_allow_html=True)

left, right = st.columns([1, 1], gap="large")

with left:
    st.markdown('<div class="card rtl">', unsafe_allow_html=True)
    st.subheader("ورودی‌ها")
    audio = st.file_uploader("فایل صوت/ویدئو مصاحبه", type=["mp3","wav","m4a","mp4","mpeg","mpga","ogg","oga","webm","flac"])
    resume = st.file_uploader("رزومه (pdf/docx/txt)", type=["pdf","docx","txt"])
    jd = st.file_uploader("آگهی شغلی (pdf/docx/txt)", type=["pdf","docx","txt"])
    st.markdown('<div class="small">نکته: متن مصاحبه خروجی خام ASR است؛ غلط‌های نوشتاری معیار قضاوت نیست.</div>', unsafe_allow_html=True)
    run = st.button("✅ تولید گزارش", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

with right:
    st.markdown('<div class="card rtl">', unsafe_allow_html=True)
    st.subheader("خروجی")
    st.markdown('<div class="small">گزارش به‌صورت HTML استاندارد نمایش داده می‌شود و Word جدول‌دار دانلود می‌شود.</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

if run:
    if not (audio and resume and jd):
        st.error("هر ۳ فایل را آپلود کن.")
        st.stop()

    progress = st.progress(0, text="شروع...")

    try:
        progress.progress(15, text="تبدیل صوت به متن...")
        interview_text = transcribe_audio_bytes(audio.getvalue(), audio.name)

        progress.progress(40, text="استخراج متن رزومه و آگهی شغلی...")
        resume_text = extract_text_from_upload(resume)
        jd_text = extract_text_from_upload(jd)

        progress.progress(70, text="تولید گزارش نهایی...")
        report_text = generate_report(resume_text, jd_text, interview_text)

        progress.progress(90, text="ساخت فایل Word جدول‌دار...")
        docx_bytes = report_to_docx_bytes(report_text)

        progress.progress(100, text="انجام شد ✅")
        st.markdown("<hr/>", unsafe_allow_html=True)

        # Display as HTML with RTL/LTR fixes + table borders
        st.markdown(markdown_to_html_with_tables(report_text), unsafe_allow_html=True)

        st.download_button(
            "⬇️ دانلود گزارش Word (جدول‌دار)",
            data=docx_bytes,
            file_name="candidate_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    except Exception as e:
        st.error(f"خطا: {e}")
